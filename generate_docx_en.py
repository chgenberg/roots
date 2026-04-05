"""Generate English version of the formulation spec as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(sh)


def _add_table(doc, headers, rows, col_widths=None, header_color="2E5E4E"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        _set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2E, 0x5E, 0x4E)

    # ── Title page ──
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("ROOTS")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x4E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Product Formulations\nShampoo  |  Conditioner  |  Body Wash")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run("Chemist Specification v1.0\n2026-03-31\nCONFIDENTIAL")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── 1. Brand profile ──
    doc.add_heading("1. Brand Profile & Requirements", level=1)

    doc.add_heading("Target Audience", level=2)
    for item in [
        "Unisex (no gendered scent profiles or product names)",
        "Sports clubs & associations \u2014 sold as a bundle (shampoo + conditioner + body wash)",
        "Age range: 16\u201365 years",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Price Point", level=2)
    for item in [
        "Consumer price: 399 SEK (~\u20ac35) for a three-product bundle",
        "COGS target: max 60\u201380 SEK per bundle (incl. packaging)",
        "Pack size: 250 ml per product (750 ml total per bundle)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Product Philosophy", level=2)
    for item in [
        "As natural as possible \u2014 not necessarily 100% organically certified",
        "Products must perform \u2014 efficacy takes priority over dogmatic naturalness",
        "Sulfate-free (no SLS/SLES)",
        "Silicone-free (no dimethicone/cyclomethicone)",
        "Paraben-free",
        "Free from MI/MCI (methylisothiazolinone / methylchloroisothiazolinone)",
        "EU-compliant (EU Cosmetic Regulation 1223/2009)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ── 2. Common design ──
    doc.add_heading("2. Shared Design Principles", level=1)

    doc.add_heading("Preservation System (all three products)", level=2)
    doc.add_paragraph(
        "System: Sodium Benzoate + Potassium Sorbate + Ethylhexylglycerin"
    )
    _add_table(doc,
        ["Ingredient", "Level", "Function"],
        [
            ["Sodium Benzoate", "0.40%", "Antibacterial, anti-yeast (pKa 4.2)"],
            ["Potassium Sorbate", "0.30%", "Anti-mold, anti-yeast (pKa 4.8)"],
            ["Ethylhexylglycerin", "0.30%", "Preservation booster, skin conditioner"],
        ],
        col_widths=[5, 2.5, 9],
    )
    p = doc.add_paragraph()
    r = p.add_run("CRITICAL: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run("Finished product pH must remain below 5.5 for this preservation system to be effective.")

    doc.add_heading("Chelator", level=2)
    doc.add_paragraph("Sodium Phytate (0.10\u20130.15%) \u2014 natural origin (rice bran), binds metal ions that destabilize preservatives and cause discoloration.")

    doc.add_heading("Fragrance Profile", level=2)
    doc.add_paragraph(
        "All three products must share a harmonized scent. Unisex, Nordic/natural character."
    )
    doc.add_paragraph('Direction: "Nordic Woods"', style="List Bullet")
    doc.add_paragraph("Top notes: Bergamot, green tea", style="List Bullet 2")
    doc.add_paragraph("Heart notes: Lavender, rosemary", style="List Bullet 2")
    doc.add_paragraph("Base notes: Cedarwood, vetiver", style="List Bullet 2")
    doc.add_paragraph(
        "Implementation: Natural fragrance composition at 0.3\u20130.5% in finished product, "
        "or a ready-made composition from Givaudan, Firmenich, or IFF with "
        '"natural origin" certification.'
    )
    doc.add_paragraph(
        "Allergen declaration: At >0.01% in rinse-off products, fragrance allergens "
        "must be declared (limonene, linalool, etc.) per EU Regulation 1223/2009."
    )

    doc.add_page_break()

    # ── 3. Shampoo ──
    doc.add_heading("3. Formulation 1: SHAMPOO", level=1)

    doc.add_heading("Design Rationale", level=2)
    doc.add_paragraph(
        "Based on ingredient analysis of 362 natural shampoos: the best-performing sulfate-free "
        "system combines Sodium Cocoyl Isethionate as a mild primary surfactant with "
        "Cocamidopropyl Betaine as co-surfactant and a glucoside surfactant for added mildness. "
        "This combination delivers rich lather, effective cleansing, and minimal irritation."
    )

    doc.add_heading("INCI Formula", level=2)
    _add_table(doc,
        ["#", "INCI Name", "Level (%)", "Function", "Phase"],
        [
            ["1", "Aqua", "to 100%", "Solvent", "A"],
            ["2", "Sodium Cocoyl Isethionate", "12.0", "Primary anionic surfactant (mild, rich lather)", "B"],
            ["3", "Cocamidopropyl Betaine", "8.0", "Co-surfactant (amphoteric, foam booster)", "A"],
            ["4", "Coco-Glucoside", "4.0", "Non-ionic co-surfactant (glucoside-based, ultra-mild)", "A"],
            ["5", "Glycerin", "3.0", "Humectant", "A"],
            ["6", "Aloe Barbadensis Leaf Juice", "2.0", "Soothing, moisturizing botanical", "A"],
            ["7", "Panthenol", "1.0", "Provitamin B5, strengthens hair shaft", "A"],
            ["8", "Guar Hydroxypropyltrimonium Chloride", "0.30", "Conditioning polymer (detangling)", "C"],
            ["9", "Niacinamide", "0.30", "Vitamin B3, follicle support, sebum balance", "C"],
            ["10", "Glyceryl Oleate", "0.50", "Re-fatting agent (reduces over-stripping)", "B"],
            ["11", "Xanthan Gum", "0.60", "Viscosity modifier (natural gum)", "C"],
            ["12", "Sodium Chloride", "0.50", "Viscosity adjustment", "C"],
            ["13", "Sodium Benzoate", "0.40", "Preservative", "C"],
            ["14", "Potassium Sorbate", "0.30", "Preservative", "C"],
            ["15", "Ethylhexylglycerin", "0.30", "Preservation booster", "C"],
            ["16", "Sodium Phytate", "0.10", "Chelator", "C"],
            ["17", "Citric Acid", "q.s.", "pH adjustment to 5.0\u20135.5", "C"],
            ["18", "Parfum (Natural)", "0.40", "Fragrance", "C"],
        ],
        col_widths=[0.8, 5.5, 2, 6, 1],
    )

    doc.add_heading("Manufacturing Instructions", level=2)
    doc.add_heading("Phase A \u2014 Water phase (room temp \u2192 40 \u00b0C)", level=3)
    for step in [
        "Weigh Aqua into the main vessel.",
        "Add Cocamidopropyl Betaine under slow agitation.",
        "Add Coco-Glucoside.",
        "Add Glycerin and Aloe Barbadensis Leaf Juice.",
        "Stir until homogeneous.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Phase B \u2014 Melt phase (65\u201370 \u00b0C)", level=3)
    for step in [
        "Melt Sodium Cocoyl Isethionate (SCI) in a separate vessel at 65\u201370 \u00b0C.",
        "Add Glyceryl Oleate to the melt.",
        "Slowly add Phase B to Phase A under stirring.",
        "Stir until SCI is fully dissolved and the blend is uniform.",
        "Allow to cool to below 40 \u00b0C.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Phase C \u2014 Cold blend (below 40 \u00b0C)", level=3)
    for step in [
        "Add Panthenol, Niacinamide, Guar Hydroxypropyltrimonium Chloride.",
        "Add Xanthan Gum (pre-disperse in glycerin if needed to avoid lumps).",
        "Add Sodium Chloride for viscosity tuning.",
        "Add preservation system.",
        "Add Sodium Phytate and Parfum.",
        "Adjust pH with Citric Acid to 5.0\u20135.5.",
        "Stir slowly for 10\u201315 minutes, avoiding air entrainment.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Specifications", level=2)
    _add_table(doc,
        ["Parameter", "Requirement"],
        [
            ["pH", "5.0\u20135.5"],
            ["Viscosity", "3,000\u20136,000 cP"],
            ["Color", "Clear to slightly opalescent"],
            ["Appearance", "Smooth gel, no visible particles"],
            ["Microbiology", "Passes EU challenge test (ISO 11930)"],
            ["Stability", "3 months accelerated (40 \u00b0C / 75% RH) with no separation"],
        ],
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # ── 4. Conditioner ──
    doc.add_heading("4. Formulation 2: CONDITIONER", level=1)

    doc.add_heading("Design Rationale", level=2)
    doc.add_paragraph(
        "Based on ingredient analysis of 179 natural conditioners: the dominant system uses "
        "Cetearyl Alcohol as the waxy base with Behentrimonium Chloride as the cationic "
        "conditioning agent. Shea butter and coconut oil are the most common natural emollients. "
        "Panthenol and hydrolyzed wheat protein are the most popular actives."
    )

    doc.add_heading("INCI Formula", level=2)
    _add_table(doc,
        ["#", "INCI Name", "Level (%)", "Function", "Phase"],
        [
            ["1", "Aqua", "to 100%", "Solvent", "A"],
            ["2", "Cetearyl Alcohol", "5.0", "Fatty co-emulsifier, viscosity, slip", "B"],
            ["3", "Behentrimonium Chloride", "2.0", "Primary cationic conditioner", "B"],
            ["4", "Glycerin", "3.0", "Humectant", "A"],
            ["5", "Butyrospermum Parkii (Shea) Butter", "2.0", "Emollient, nourishing", "B"],
            ["6", "Cocos Nucifera (Coconut) Oil", "1.5", "Emollient, penetrates hair shaft", "B"],
            ["7", "Aloe Barbadensis Leaf Juice", "2.0", "Soothing, moisturizing", "A"],
            ["8", "Panthenol", "1.0", "Provitamin B5, strengthens hair", "C"],
            ["9", "Hydrolyzed Wheat Protein", "0.50", "Protein repair, volume, strength", "C"],
            ["10", "Argania Spinosa Kernel Oil", "0.50", "Argan oil, shine, frizz control", "B"],
            ["11", "Stearamidopropyl Dimethylamine", "1.0", "Secondary conditioner (detangling, softness)", "B"],
            ["12", "Guar Hydroxypropyltrimonium Chloride", "0.20", "Conditioning polymer", "C"],
            ["13", "Lactic Acid", "0.30", "Cuticle smoothing, shine enhancement", "C"],
            ["14", "Tocopherol", "0.10", "Vitamin E, antioxidant (protects oils)", "B"],
            ["15", "Sodium Benzoate", "0.40", "Preservative", "C"],
            ["16", "Potassium Sorbate", "0.30", "Preservative", "C"],
            ["17", "Ethylhexylglycerin", "0.30", "Preservation booster", "C"],
            ["18", "Sodium Phytate", "0.10", "Chelator", "C"],
            ["19", "Citric Acid", "q.s.", "pH adjustment to 4.5\u20135.0", "C"],
            ["20", "Parfum (Natural)", "0.40", "Fragrance", "C"],
        ],
        col_widths=[0.8, 5.5, 2, 6, 1],
    )

    doc.add_heading("Manufacturing Instructions", level=2)
    doc.add_heading("Phase A \u2014 Water phase (70\u201375 \u00b0C)", level=3)
    for step in [
        "Weigh Aqua, Glycerin, and Aloe Barbadensis Leaf Juice into the main vessel.",
        "Heat to 70\u201375 \u00b0C under agitation.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Phase B \u2014 Oil phase (70\u201375 \u00b0C)", level=3)
    for step in [
        "In a separate vessel: melt Cetearyl Alcohol at ~70 \u00b0C.",
        "Add Behentrimonium Chloride, Stearamidopropyl Dimethylamine, Shea Butter, Coconut Oil, Argan Oil.",
        "Add Tocopherol.",
        "Stir until fully melted and homogeneous.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Emulsification", level=3)
    for step in [
        "Slowly pour Phase B into Phase A under high-shear mixing.",
        "Emulsify for 3\u20135 minutes until smooth and homogeneous.",
        "Continue stirring while cooling to ~40 \u00b0C.",
    ]:
        doc.add_paragraph(step, style="List Number")

    p = doc.add_paragraph()
    r = p.add_run("NOTE: ")
    r.bold = True
    p.add_run("Cooling too rapidly will produce a grainy texture. Allow gradual cooling.")

    doc.add_heading("Phase C \u2014 Cold blend (below 40 \u00b0C)", level=3)
    for step in [
        "Add Panthenol, Hydrolyzed Wheat Protein, Guar Hydroxypropyltrimonium Chloride.",
        "Add preservation system.",
        "Add Sodium Phytate and Parfum.",
        "Adjust pH with Citric Acid / Lactic Acid to 4.5\u20135.0.",
        "Stir slowly for an additional 15\u201320 minutes.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Specifications", level=2)
    _add_table(doc,
        ["Parameter", "Requirement"],
        [
            ["pH", "4.5\u20135.0"],
            ["Viscosity", "15,000\u201330,000 cP"],
            ["Color", "White to off-white"],
            ["Appearance", "Smooth creamy emulsion, no graininess"],
            ["Microbiology", "Passes EU challenge test (ISO 11930)"],
            ["Stability", "3 months accelerated with no separation or viscosity loss"],
        ],
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # ── 5. Body Wash ──
    doc.add_heading("5. Formulation 3: BODY WASH", level=1)

    doc.add_heading("Design Rationale", level=2)
    doc.add_paragraph(
        "Based on ingredient analysis of 145 body wash products: the mildest and highest-rated "
        "formulas use Cocamidopropyl Betaine as the primary surfactant, complemented by "
        "Coco-Glucoside and Decyl Glucoside for the gentlest possible cleansing profile. "
        "A higher level of glycerin (5%) and aloe vera provides moisture balance. No sulfates."
    )

    doc.add_heading("INCI Formula", level=2)
    _add_table(doc,
        ["#", "INCI Name", "Level (%)", "Function", "Phase"],
        [
            ["1", "Aqua", "to 100%", "Solvent", "A"],
            ["2", "Cocamidopropyl Betaine", "10.0", "Primary amphoteric surfactant (extremely mild)", "A"],
            ["3", "Coco-Glucoside", "5.0", "Co-surfactant (non-ionic, plant-derived)", "A"],
            ["4", "Decyl Glucoside", "3.0", "Co-surfactant (added mildness, foam stabilizer)", "A"],
            ["5", "Glycerin", "5.0", "Humectant (higher level for skin care)", "A"],
            ["6", "Aloe Barbadensis Leaf Juice", "2.0", "Soothing, anti-inflammatory", "A"],
            ["7", "Sodium Lauroyl Sarcosinate", "2.0", "Mild anionic surfactant (lather boost)", "A"],
            ["8", "Panthenol", "0.50", "Provitamin B5, skin barrier support", "B"],
            ["9", "Niacinamide", "0.30", "Vitamin B3, skin function", "B"],
            ["10", "Chamomilla Recutita Flower Extract", "0.30", "Chamomile, soothing, anti-inflammatory", "B"],
            ["11", "Glyceryl Oleate", "0.50", "Re-fatting agent", "A"],
            ["12", "Sodium Cocoamphoacetate", "2.0", "Amphoteric co-surfactant", "A"],
            ["13", "Hydroxypropyl Methylcellulose", "0.40", "Viscosity modifier", "B"],
            ["14", "Xanthan Gum", "0.30", "Viscosity modifier", "B"],
            ["15", "Sodium Benzoate", "0.40", "Preservative", "B"],
            ["16", "Potassium Sorbate", "0.30", "Preservative", "B"],
            ["17", "Ethylhexylglycerin", "0.30", "Preservation booster", "B"],
            ["18", "Sodium Phytate", "0.10", "Chelator", "B"],
            ["19", "Citric Acid", "q.s.", "pH adjustment to 5.0\u20135.5", "B"],
            ["20", "Parfum (Natural)", "0.40", "Fragrance", "B"],
        ],
        col_widths=[0.8, 5.5, 2, 6, 1],
    )

    doc.add_heading("Manufacturing Instructions", level=2)
    doc.add_heading("Phase A \u2014 Main phase (room temp \u2192 35 \u00b0C)", level=3)
    for step in [
        "Weigh Aqua into the main vessel.",
        "Slowly add Cocamidopropyl Betaine under gentle agitation.",
        "Add Coco-Glucoside and Decyl Glucoside.",
        "Add Sodium Lauroyl Sarcosinate and Sodium Cocoamphoacetate.",
        "Add Glycerin, Aloe Barbadensis Leaf Juice, Glyceryl Oleate.",
        "Stir until homogeneous. Avoid air entrainment.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Phase B \u2014 Additives (below 35 \u00b0C)", level=3)
    for step in [
        "Pre-disperse HPMC and Xanthan Gum in a small amount of glycerin, then add to batch.",
        "Add Panthenol, Niacinamide, Chamomilla Recutita extract.",
        "Add preservation system.",
        "Add Sodium Phytate and Parfum.",
        "Adjust pH with Citric Acid to 5.0\u20135.5.",
        "Stir slowly for 10\u201315 minutes.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Specifications", level=2)
    _add_table(doc,
        ["Parameter", "Requirement"],
        [
            ["pH", "5.0\u20135.5"],
            ["Viscosity", "2,000\u20135,000 cP"],
            ["Color", "Clear to slightly opalescent"],
            ["Appearance", "Smooth gel"],
            ["Microbiology", "Passes EU challenge test (ISO 11930)"],
            ["Stability", "3 months accelerated with no separation"],
        ],
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # ── 6. Shared overview ──
    doc.add_heading("6. Shared Ingredient Overview", level=1)
    doc.add_paragraph("Ingredients common to all three products, creating a cohesive product family:")
    _add_table(doc,
        ["Ingredient", "Shampoo", "Conditioner", "Body Wash", "Purpose"],
        [
            ["Aqua", "base", "base", "base", "Solvent"],
            ["Glycerin", "3%", "3%", "5%", "Humectant"],
            ["Aloe Barbadensis Leaf Juice", "2%", "2%", "2%", "Soothing botanical"],
            ["Panthenol", "1%", "1%", "0.5%", "Provitamin B5"],
            ["Sodium Benzoate", "0.4%", "0.4%", "0.4%", "Preservative"],
            ["Potassium Sorbate", "0.3%", "0.3%", "0.3%", "Preservative"],
            ["Ethylhexylglycerin", "0.3%", "0.3%", "0.3%", "Preservation booster"],
            ["Sodium Phytate", "0.1%", "0.1%", "0.1%", "Chelator"],
            ["Parfum (Natural)", "0.4%", "0.4%", "0.4%", "Fragrance"],
        ],
        col_widths=[5, 2, 2, 2, 4],
    )

    # ── 7. Suppliers ──
    doc.add_heading("7. Suggested Raw Material Suppliers", level=1)
    _add_table(doc,
        ["Raw Material", "Supplier Suggestions"],
        [
            ["Sodium Cocoyl Isethionate", "BASF (Plantapon SCI), Innospec"],
            ["Cocamidopropyl Betaine", "Evonik (Tego Betain), Solvay"],
            ["Coco-Glucoside / Decyl Glucoside", "BASF (Plantacare), Galaxy Surfactants"],
            ["Behentrimonium Chloride", "Evonik (Varisoft BT 85), Clariant"],
            ["Cetearyl Alcohol", "BASF, KLK Oleo"],
            ["Shea Butter", "AAK (Lipex SheaSoft)"],
            ["Preservation System", "Lonza (Geogard Ultra), Ashland"],
            ["Fragrance", "Givaudan, Firmenich, IFF"],
        ],
        col_widths=[6, 10],
    )

    doc.add_page_break()

    # ── 8. Testing ──
    doc.add_heading("8. Testing Protocol", level=1)

    doc.add_heading("Phase 1: Lab Scale (1\u20132 kg batches)", level=2)
    for s in [
        "Formulate all three products.",
        "Measure pH, viscosity, appearance.",
        "Basic foam test (cylinder method).",
        "Sensory evaluation (scent, texture, lathering, rinsability).",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Phase 2: Stability Testing", level=2)
    for s in [
        "Centrifuge test: 3,000 rpm, 30 min \u2014 no separation permitted.",
        "Freeze-thaw cycling: 5 cycles (\u221210 \u00b0C / +40 \u00b0C, 24 h each).",
        "Accelerated aging: 40 \u00b0C / 75% RH for 3 months.",
        "Room temperature: 25 \u00b0C for 6 months (long-term stability).",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Phase 3: Microbiological Testing", level=2)
    for s in [
        "Challenge test (ISO 11930 / EP Method A) \u2014 mandatory for the EU market.",
        "Total aerobic microbial count at batch release.",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Phase 4: Consumer Panel", level=2)
    for s in [
        "Panel test with 20\u201330 participants (unisex, varied hair types).",
        "Assess: foam quality, cleansing, hair feel after drying, skin feel, scent.",
        "Benchmark against leading natural brands (Maria Nila, Sachajuan, Davines).",
    ]:
        doc.add_paragraph(s, style="List Number")

    # ── 9. Regulatory ──
    doc.add_heading("9. Regulatory Compliance", level=1)

    doc.add_heading("EU Cosmetic Regulation 1223/2009", level=2)
    for s in [
        "All ingredients must be listed in the EU CosIng database.",
        "CPNP notification required before placing the product on the market.",
        "A Responsible Person (RP) within the EU must be designated.",
        "A Product Information File (PIF) with safety assessment is required.",
        "All fragrance allergens above 0.01% in rinse-off products must be declared.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Permitted Claims", level=2)
    doc.add_paragraph("These formulations support the following claims:")
    for s in [
        '"Sulfate-free"',
        '"Silicone-free"',
        '"Paraben-free"',
        '"Vegan" (verify with each supplier that no animal-derived raw materials are used)',
        '"With natural ingredients"',
        '"Mild / Gentle"',
    ]:
        doc.add_paragraph(s, style="List Bullet")

    p = doc.add_paragraph()
    r = p.add_run("Avoid: ")
    r.bold = True
    p.add_run('"Natural" as a standalone claim without certification \u2014 may be considered misleading.')

    # ── 10. Cost ──
    doc.add_heading("10. Cost Estimate", level=1)
    _add_table(doc,
        ["Component", "Shampoo", "Conditioner", "Body Wash"],
        [
            ["Raw materials (250 ml)", "8\u201312 SEK", "10\u201315 SEK", "8\u201312 SEK"],
            ["Packaging (bottle + cap)", "5\u20138 SEK", "5\u20138 SEK", "5\u20138 SEK"],
            ["Manufacturing + filling", "3\u20135 SEK", "3\u20135 SEK", "3\u20135 SEK"],
            ["Total per product", "16\u201325 SEK", "18\u201328 SEK", "16\u201325 SEK"],
            ["Bundle (3 pcs)", "", "50\u201378 SEK", ""],
        ],
        col_widths=[5, 3.5, 3.5, 3.5],
    )

    doc.add_paragraph(
        "At a consumer price of 399 SEK and an assumed 50% discount to clubs (wholesale 200 SEK): "
        "gross margin of 60\u201375% at contract manufacturing volumes."
    )
    doc.add_paragraph(
        "Prices are based on Northern European supplier pricing at MOQ 500\u20131,000 kg. "
        "Unit costs decrease at higher volumes."
    )

    # ── Save ──
    path = "Roots_Formulations_Chemist_Spec_EN.docx"
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build_document()
