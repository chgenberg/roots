#!/usr/bin/env python3
"""Bygger Roots säljguide (.docx) för säljare."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
WEB = os.path.abspath(os.path.join(ROOT, "..", "..", "apps", "web", "public"))
ASSETS = os.path.join(ROOT, "assets")
OUT = os.path.join(ROOT, "Roots_Saljguide.docx")

INK = RGBColor(0x1D, 0x1D, 0x1B)
SAND_DARK = RGBColor(0x7F, 0x71, 0x5B)
FOREST = RGBColor(0x6B, 0x79, 0x4F)
TERRA = RGBColor(0xE1, 0x87, 0x54)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SAND_LIGHT = RGBColor(0xD5, 0xCA, 0xBF)
OFFWHITE = RGBColor(0xFA, 0xF6, 0xEF)

HEAD = "Alan Sans"
BODY = "Inter 18pt"


def set_font(run, name=BODY, size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)


def shade(paragraph, hexcolor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def border(paragraph, color="6B794F", size=18, side="left", space=14):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pbdr.append(el)
    pPr.append(pbdr)


def space_after(p, pts):
    p.paragraph_format.space_after = Pt(pts)


def space_before(p, pts):
    p.paragraph_format.space_before = Pt(pts)


def H1(doc, text, color=INK):
    p = doc.add_paragraph()
    space_before(p, 18); space_after(p, 6)
    r = p.add_run(text)
    set_font(r, HEAD, 21, color, bold=True)
    return p


def H2(doc, text, color=FOREST):
    p = doc.add_paragraph()
    space_before(p, 14); space_after(p, 4)
    r = p.add_run(text)
    set_font(r, HEAD, 14.5, color, bold=True)
    return p


def kicker(doc, text, color=SAND_DARK):
    p = doc.add_paragraph()
    space_after(p, 2)
    r = p.add_run(text.upper())
    set_font(r, BODY, 9, color, bold=True)
    rpr = r._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "40"); rpr.append(sp)
    return p


def body(doc, text, size=11, color=INK, after=8, bold=False, italic=False):
    p = doc.add_paragraph()
    space_after(p, after)
    p.paragraph_format.line_spacing = 1.32
    r = p.add_run(text)
    set_font(r, BODY, size, color, bold=bold, italic=italic)
    return p


def bullet(doc, text, color=INK, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    space_after(p, 3)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead:
        r = p.add_run(bold_lead + "  ")
        set_font(r, BODY, 11, INK, bold=True)
    r = p.add_run(text)
    set_font(r, BODY, 11, color)
    return p


def callout(doc, title, text, fill="F1EBE2", barcolor="6B794F", title_color="6B794F"):
    p = doc.add_paragraph()
    space_before(p, 6); space_after(p, 0)
    shade(p, fill); border(p, barcolor, 24, "left", 16)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(title)
    set_font(r, HEAD, 12, RGBColor.from_string(title_color), bold=True)
    p2 = doc.add_paragraph()
    shade(p2, fill); border(p2, barcolor, 24, "left", 16)
    p2.paragraph_format.left_indent = Pt(12)
    p2.paragraph_format.right_indent = Pt(8)
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.line_spacing = 1.3
    r = p2.add_run(text)
    set_font(r, BODY, 10.5, INK)
    return p2


def img(doc, name, width_in=6.3):
    path = os.path.join(ASSETS, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def web_img(doc, relpath, width_in=6.3):
    path = os.path.join(WEB, relpath)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def rule(doc, color="E5DDD2"):
    p = doc.add_paragraph()
    space_before(p, 4); space_after(p, 4)
    border(p, color, 8, "bottom", 1)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, HEAD, 10.5, WHITE, bold=True)
        shade(p, "6B794F")
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "6B794F")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(val)
            set_font(r, BODY, 9.8, INK, bold=(i == 0))
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    # Base style
    style = doc.styles["Normal"]
    style.font.name = BODY
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    for sec in doc.sections:
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # ───────── COVER ─────────
    web_img(doc, "brand/roots-logo-black.png", width_in=2.6)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Säljguide")
    set_font(r, HEAD, 40, INK, bold=True)
    p = doc.add_paragraph(); space_after(p, 2)
    r = p.add_run("Riktigt bra hårvård från Norden \u2013\nsom hjälper svenskt föreningsliv.")
    set_font(r, HEAD, 18, FOREST, bold=True)
    doc.add_paragraph()
    body(doc, "Det här är ditt fickbibliotek när du säljer Roots. Här finns berättelsen, "
              "vetenskapen, ingredienserna och de konkreta orden du behöver \u2013 utan krångel "
              "och utan överdrifter. Läs igenom en gång, återkom till det du behöver.", 12, SAND_DARK)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Internt säljmaterial  \u00b7  Version 1.0  \u00b7  2026")
    set_font(r, BODY, 9, SAND_DARK)
    page_break(doc)

    build_intro(doc)
    build_products(doc)
    build_science(doc)
    build_ingredients(doc)
    build_talk(doc)
    build_market(doc)
    build_howto(doc)
    build_faq(doc)
    build_sources(doc)

    doc.save(OUT)
    print("Saved", os.path.relpath(OUT, ROOT))


def build_intro(doc):
    kicker(doc, "Varför vi gör det här")
    H1(doc, "Två saker på en gång: bättre produkter och starkare föreningsliv")
    body(doc, "Svenskt föreningsliv går ofta runt på försäljning. Strumpor, kataloger, "
              "godis, toapapper, rabatthäften. Det funkar \u2013 men det är sällan något man längtar "
              "efter att köpa, och nästan aldrig något man köper igen. Pengarna kommer in en gång, "
              "sen är det slut.")
    body(doc, "Roots vänder på det. Vi säljer hudvård och hårvård i premiumklass som folk faktiskt "
              "vill ha, använder varje dag och kommer tillbaka till. Det betyder att en försäljning "
              "kan bli till återkommande intäkter \u2013 samtidigt som kunden får en produkt som är "
              "snällare mot både håret och planeten.")
    callout(doc, "Det här är löftet",
            "Vi vill bevisa att man kan sälja riktigt bra saker OCH göra gott. Varje flaska som säljs "
            "ger pengar till en förening i Sverige \u2013 och en bättre dag för någons hår. Inga genvägar "
            "på kvaliteten, inga överdrifter i löftena.")
    body(doc, "Som säljare är du inte en katalogutdelare. Du är ansiktet för ett nordiskt varumärke "
              "som bryr sig. Det är därför den här guiden finns: så att du kan berätta varför \u2013 "
              "enkelt, ärligt och med stolthet.", after=10)
    page_break(doc)


def build_products(doc):
    kicker(doc, "Sortimentet")
    H1(doc, "Tre formuleringar, en filosofi")
    body(doc, "Hela serien är byggd kring samma idé: rengör mjukt, lugna hudbotten och lämna "
              "håret starkare än innan. Sulfatsnålt, biologiskt nedbrytbara mjukgörare och aktiva "
              "ingredienser med dokumenterad effekt.")
    img(doc, "product_line.png", 6.4)
    body(doc, "")
    H2(doc, "Pure Shampoo \u00b7 Syricalm")
    body(doc, "En mild schampotvätt som rengör utan att torka ut. Syricalm lugnar en känslig "
              "hudbotten \u2013 perfekt för den som lätt blir kliande eller irriterad.", after=4)
    H2(doc, "Pure Conditioner (less perfume) \u00b7 Syricalm")
    body(doc, "Balsam med medvetet låg parfymhalt för den som vill ha minimal doft. Reder ut, "
              "återfuktar och gör håret lättkammat. Mjukgöraren är en modern, biologiskt nedbrytbar "
              "esterquat \u2013 skonsam mot både hår och avlopp.", after=4)
    H2(doc, "Body Wash (more perfume) \u00b7 Syricalm")
    body(doc, "En krämig, lite mer doftande duschtvätt med milda tensider (cocamidopropyl betaine "
              "och sodium lauroyl sarcosinate). Rengör skönt och lämnar huden mjuk, inte stram.", after=4)
    callout(doc, "På gång \u2013 ny lansering att teasa",
            "En Body Wash med 0,25 % mentol och Multimoist (rödbeta + prebiotika) testas just nu. "
            "Den ger en pigg, sval känsla i huden. Nämn den gärna som \u201ekommer snart\u201d för nyfikna kunder \u2013 "
            "men lova inget datum.",
            fill="EDF1E9", barcolor="6B794F")
    page_break(doc)
def build_science(doc):
    kicker(doc, "Det som gör Roots annorlunda")
    H1(doc, "Vi jobbar MED hudbotten \u2013 inte emot den")
    body(doc, "Det här är vår hemliga sås, förklarad så att vem som helst förstår. Håret växer "
              "ur hudbotten, och hudbotten är levande. Två små system håller den i balans:")
    img(doc, "ecs_microbiome.png", 6.4)
    body(doc, "")
    H2(doc, "1. Hudbottens eget balanssystem (ECS)")
    body(doc, "Forskningen har visat att huden \u2013 inklusive hårfolliklar och talgkörtlar \u2013 har ett "
              "eget reglersystem som kallas endocannabinoidsystemet. Forskare kallar det till och med "
              "hudens \u201ec(ut)annabinoid\u201d-system. Dess uppgift är att hålla balansen: lagom med talg, "
              "lugn hud, en barriär som mår bra. När balansen rubbas blir det lätt kliande, torrt och känsligt.")
    H2(doc, "2. Hudbottens mikrobiom")
    body(doc, "På hudbotten lever ett helt ekosystem av bakterier och jäst. När de är i balans håller "
              "de varandra i schack och hjälper huden att må bra. Hård tvätt och aggressiva tensider kan "
              "rubba balansen \u2013 därför sitter mjäll och irritation så ofta i andra änden av en för stark schampoflaska.")
    callout(doc, "Så här säger vi det till kunden",
            "\u201eHåret och hudbotten har faktiskt ett eget balanssystem och ett levande mikrobiom. "
            "Roots är gjort för att jobba med den naturliga balansen \u2013 lugna, återfukta och nära \u2013 "
            "istället för att tvätta bort den. Friskt hår börjar i en hudbotten i balans.\u201c")
    body(doc, "På så sätt blir vetenskapen inte skrytig, utan en omtanke: vi tar hand om grunden "
              "som håret växer ur.", after=10)
    img(doc, "science_balance.png", 6.2)
    page_break(doc)


def build_ingredients(doc):
    kicker(doc, "Nyckelingredienser")
    H1(doc, "Få ingredienser \u2013 men de som finns gör jobbet")
    body(doc, "Du behöver inte kunna hela INCI-listan utantill. Lär dig de fyra hjältarna nedan, "
              "så kan du svara på nästan vad som helst.")
    img(doc, "ingredient_cards.png", 6.4)
    body(doc, "")
    add_table(doc,
        ["Ingrediens", "Vad det är", "Varför det spelar roll"],
        [
            ["SyriCalm\u2122", "Extrakt av vass (Phragmites) och svampen Poria cocos.",
             "Lugnar kliande och känslig hudbotten, dämpar rodnad och hjälper hudbarriären att återhämta sig. NATRUE- och COSMOS-godkänd."],
            ["MultiMoist\u2122", "Prebiotiska fruktosockerarter (FOS) och rödbetsextrakt.",
             "Binder fukt inne i hårstrået \u2013 mindre hårbrott, frizz och statik. Skyddar färgen och närar hudbottens mikrobiom."],
            ["Panthenol", "Provitamin B5.",
             "Klassisk fukt- och konditioneringsingrediens. Gör håret mjukt, spänstigt och lättkammat."],
            ["Mentol", "Renframställd svalkande botanik.",
             "Ger en pigg, sval känsla i hudbotten. Doseras lågt (0,25 %) \u2013 en njutning, inte en chock."],
            ["Tocopherol", "E-vitamin (antioxidant).",
             "Skyddar formuleringen och håret mot oxidation."],
            ["Milda tensider", "Cocamidopropyl betaine, sodium lauroyl sarcosinate.",
             "Rengör skonsamt utan att strippa huden \u2013 grunden i en sulfatsnål tvätt."],
        ],
        widths=[1.3, 2.1, 3.0])
    body(doc, "")
    callout(doc, "Bra att veta",
            "Både SyriCalm och MultiMoist kommer från tyska CLR Berlin och är certifierade enligt "
            "NATRUE och COSMOS \u2013 oberoende standarder för naturlig kosmetik. Det är inte marknadsföring, "
            "det är tredjepartskontroll.",
            fill="F1EBE2", barcolor="7F715B", title_color="7F715B")
    page_break(doc)


def build_talk(doc):
    kicker(doc, "Så pratar du om det")
    H1(doc, "Säg det starkt \u2013 men aldrig mer än vi får")
    body(doc, "Roots är kosmetik, inte medicin. Vi får prata om hur produkten känns, vårdar och "
              "balanserar \u2013 men vi lovar aldrig att den botar något eller får hår att växa. Det är inte "
              "bara regler; det är trovärdighet. Vi vinner på att vara ärliga.")
    add_table(doc,
        ["Säg gärna", "Undvik"],
        [
            ["\u201eLugnar och balanserar hudbotten.\u201c", "\u201eBotar mjäll eller eksem.\u201c"],
            ["\u201eGör håret mjukt, starkare och lättkammat.\u201c", "\u201eFår håret att växa / stoppar håravfall.\u201c"],
            ["\u201eÅterfuktar och skyddar mot frizz.\u201c", "\u201eMedicinsk effekt eller behandling.\u201c"],
            ["\u201eStödjer hudens barriär och mikrobiom.\u201c", "\u201eDödar bakterier / antibakteriellt.\u201c"],
            ["\u201eSnällt mot huden och naturen.\u201c", "\u201e100 % naturligt\u201c (om det inte stämmer)."],
        ],
        widths=[3.2, 3.2])
    body(doc, "")
    callout(doc, "Tumregeln",
            "Prata om känslan och omtanken, inte om diagnoser. \u201eDet här gör hudbotten lugn och "
            "håret glatt\u201c slår alltid \u201edet botar X\u201c \u2013 både juridiskt och i kundens hjärta.",
            fill="FBEFE9", barcolor="E18754", title_color="C76A50")
    page_break(doc)


def build_market(doc):
    kicker(doc, "Marknaden och varför nu")
    H1(doc, "Föreningssverige har gått digitalt \u2013 men sortimentet har inte följt med")
    body(doc, "Föreningsförsäljning är en jätteindustri. Newbody har ensamt hjälpt föreningar att "
              "samla in över 1,7 miljarder kronor på snart 40 år, och nya digitala plattformar växer "
              "snabbt med personliga säljlänkar, QR-koder och betalning via Klarna och Swish.")
    body(doc, "Men titta på VAD som säljs: strumpor, godis, kaffe, toapapper, rabatthäften. "
              "Lågt engagemang, ofta ett engångsköp. Ingen återkommer för att köpa mer toapapper av "
              "kärlek till varumärket. Det är precis där Roots har en öppen flank.")
    img(doc, "positioning_map.png", 5.4)
    body(doc, "")
    H2(doc, "Tre saker som gör att premium vinner i kassan")
    bullet(doc, "En riktigt bra produkt har ett högre värde \u2013 större ordrar, mer per sälj.", bold_lead="Högre ordervärde.")
    bullet(doc, "Hudvård tar slut. När flaskan är tom vill kunden ha en ny \u2013 det öppnar för återköp och prenumeration.", bold_lead="Återkommande intäkter.")
    bullet(doc, "Det är roligare att sälja något man själv är stolt över. Stolthet syns och smittar.", bold_lead="Lätt att sälja.")
    img(doc, "value_compare.png", 5.8)
    body(doc, "Siffrorna ovan är illustrativa för att visa principen: även med en rimlig marginal "
              "ger ett högre ordervärde och återköp mer till föreningen över tid än ett billigt engångsköp.", 9.5, SAND_DARK)
    page_break(doc)


def build_howto(doc):
    kicker(doc, "Så säljer du")
    H1(doc, "Fyra steg \u2013 och din länk gör jobbet")
    img(doc, "sales_journey.png", 6.4)
    body(doc, "")
    H2(doc, "Dina bästa knep")
    bullet(doc, "Berätta varför INNAN du berättar vad. Människor köper berättelsen: nordiskt, snällt mot håret, stöttar vår förening.")
    bullet(doc, "Gör det personligt: \u201eJag säljer för [föreningen] \u2013 varje flaska hjälper oss att [målet].\u201c")
    bullet(doc, "Använd QR-koden vid träningar, matcher och på jobbet. En skärm räcker.")
    bullet(doc, "Följ upp en gång. \u201eHann du titta på länken?\u201c sluter fler affärer än något annat.")
    bullet(doc, "Tänk återköp: \u201eNär flaskan tar slut finns vi kvar.\u201c")
    body(doc, "")
    H2(doc, "30-sekunderspitchen (lär dig utantill)")
    callout(doc, "Säg ungefär så här",
            "\u201eVi säljer Roots \u2013 nordisk hår- och hudvård i premiumklass. Det är gjort för att lugna "
            "hudbotten och stärka håret på riktigt, utan onödiga tillsatser, och det är certifierat naturligt. "
            "Varje flaska du köper ger pengar till [föreningen]. Det är bra för ditt hår och bra för oss \u2013 "
            "får jag skicka dig länken?\u201c",
            fill="EDF1E9", barcolor="6B794F")
    page_break(doc)


def build_faq(doc):
    kicker(doc, "Vanliga frågor")
    H1(doc, "Så svarar du när någon tvekar")
    qa = [
        ("\u201eDet låter dyrt.\u201c",
         "\u201eDet kostar som ett bra schampo i butik \u2013 men här går en del till föreningen, och kvaliteten "
         "är i premiumklass. Du betalar för något du använder varje dag.\u201c"),
        ("\u201eFungerar det verkligen?\u201c",
         "\u201eDe aktiva ingredienserna är dokumenterade och certifierade (NATRUE/COSMOS). Vi lovar inga mirakel \u2013 "
         "men det är gjort för att lugna hudbotten och göra håret mjukt och starkt. Många märker skillnad på känslan direkt.\u201c"),
        ("\u201eJag har känslig hudbotten.\u201c",
         "\u201eDå kan det här passa extra bra \u2013 SyriCalm finns just för att lugna känslig, kliande hudbotten. "
         "Det är sulfatsnålt och skonsamt.\u201c"),
        ("\u201eVad är det med endocannabinoid och mikrobiom?\u201c",
         "\u201eHudbotten har ett eget balanssystem och ett levande mikrobiom. Roots är gjort för att jobba "
         "med den balansen istället för att tvätta bort den. Friskt hår börjar i en lugn hudbotten.\u201c"),
        ("\u201eÄr det miljövänligt?\u201c",
         "\u201eVi använder biologiskt nedbrytbara mjukgörare, sulfatsnåla tvättar och certifierat naturliga aktiver. "
         "Det är medvetna val hela vägen.\u201c"),
        ("\u201eJag köper redan mitt schampo.\u201c",
         "\u201eHelt klart \u2013 men nästa gång det tar slut, varför inte testa något som är snällare mot håret OCH "
         "hjälper föreningen? Jag skickar länken så har du den när det blir dags.\u201c"),
    ]
    for q, a in qa:
        p = doc.add_paragraph(); space_after(p, 1)
        r = p.add_run(q); set_font(r, HEAD, 11.5, INK, bold=True)
        body(doc, a, 10.8, SAND_DARK, after=8)
    page_break(doc)


def build_sources(doc):
    kicker(doc, "För den nyfikne")
    H1(doc, "Källor och underlag")
    body(doc, "Du behöver aldrig citera det här för en kund \u2013 men det är bra att veta att det finns på riktigt.", 10.5, SAND_DARK)
    H2(doc, "Ingredienser")
    body(doc, "CLR Berlin \u2013 SyriCalm\u2122 CLR (PC): produktblad, Phragmites communis + Poria cocos, lugnande/barriär, NATRUE/COSMOS.", 9.8, INK, after=3)
    body(doc, "CLR Berlin \u2013 MultiMoist\u2122 CLR: produktblad, fruktooligosackarider + Beta vulgaris, fukt/hår/prebiotisk mikrobiom-effekt.", 9.8, INK, after=3)
    H2(doc, "Vetenskap")
    body(doc, "Toth et al., \u201eCannabinoid Signaling in the Skin\u201c (PMC6429381) \u2013 hudens endocannabinoidsystem (\u201ec(ut)annabinoid\u201c).", 9.8, INK, after=3)
    body(doc, "Biro et al., \u201eThe endocannabinoid system of the skin\u201c (PMC2757311) \u2013 ECS reglerar balans i hud och hårfollikel.", 9.8, INK, after=3)
    body(doc, "Decoding scalp health and microbiome dysbiosis in dandruff (bioRxiv 2024) \u2013 mikrobiomets roll för hudbottens hälsa.", 9.8, INK, after=3)
    H2(doc, "Marknad")
    body(doc, "Newbody Family, Klubbförsäljning.se, Föreningshäftet, Godio \u2013 publika uppgifter om upplägg, marginaler och intäkter (2025).", 9.8, INK, after=3)
    body(doc, "")
    rule(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Roots \u00b7 Riktigt bra produkter. Starkare föreningsliv.")
    set_font(r, HEAD, 12, FOREST, bold=True)


build()

