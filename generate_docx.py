"""Generate a professional Word document from the formulation spec."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(sh)


def _add_table(doc, headers, rows, col_widths=None, header_color="2E5E4E"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        _set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x2E, 0x5E, 0x4E)

    # ── Title page ──
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("ROOTS")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x4E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Produktformuleringar\nSchampo  |  Balsam  |  Body Wash")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run("Kemistspecifikation v1.0\n2026-03-31\nKONFIDENTIELLT")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── 1. Brand profile ──
    doc.add_heading("1. Varumärkesprofil & Kravspecifikation", level=1)

    doc.add_heading("Målgrupp", level=2)
    for item in [
        "Unisex (inga könade doftprofiler eller produktnamn)",
        "Föreningsliv/klubbar — säljs som paket (schampo + balsam + body wash)",
        "Åldersspann: 16–65 år",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Prisklass", level=2)
    for item in [
        "Konsumentpris: 399 SEK för paket om tre produkter",
        "COGS-mål: max 60–80 SEK per paket (inkl. förpackning)",
        "Förpackning: 250 ml per produkt (totalt 750 ml per paket)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Produktfilosofi", level=2)
    for item in [
        "Så naturligt som möjligt — inte nödvändigtvis 100 % ekologiskt certifierat",
        "Produkterna ska fungera — prestandakrav går före dogmatisk naturlighet",
        "Sulfatfri (inga SLS/SLES)",
        "Silikonfri (inga dimethicone/cyclomethicone)",
        "Parabensfri",
        "Fri från MI/MCI (methylisothiazolinone / methylchloroisothiazolinone)",
        "EU-compliant (EU Cosmetic Regulation 1223/2009)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ── 2. Common design ──
    doc.add_heading("2. Gemensamma designprinciper", level=1)

    doc.add_heading("Konserveringssystem (alla tre produkter)", level=2)
    doc.add_paragraph(
        "System: Sodium Benzoate + Potassium Sorbate + Ethylhexylglycerin"
    )
    _add_table(doc,
        ["Ingrediens", "Andel", "Funktion"],
        [
            ["Sodium Benzoate", "0.40%", "Antibakteriell, anti-jäst (pKa 4.2)"],
            ["Potassium Sorbate", "0.30%", "Anti-mögel, anti-jäst (pKa 4.8)"],
            ["Ethylhexylglycerin", "0.30%", "Konserveringsförstärkare, hudkonditionerare"],
        ],
        col_widths=[5, 2.5, 9],
    )
    p = doc.add_paragraph()
    r = p.add_run("KRITISKT: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run("Slutproduktens pH måste ligga under 5.5 för att konserveringssystemet ska vara effektivt.")

    doc.add_heading("Kelator", level=2)
    doc.add_paragraph("Sodium Phytate (0.10–0.15%) — naturligt ursprung (ris), binder metalljoner.")

    doc.add_heading("Doftprofil", level=2)
    doc.add_paragraph(
        "Alla tre produkter ska harmoniera i doft. Unisex, nordisk/naturlig känsla."
    )
    doc.add_paragraph('Riktning: "Nordic Woods"', style="List Bullet")
    doc.add_paragraph("Topptoner: Bergamott, grönt te", style="List Bullet 2")
    doc.add_paragraph("Hjärttoner: Lavendel, rosmarin", style="List Bullet 2")
    doc.add_paragraph("Bastoner: Cederträ, vetiver", style="List Bullet 2")
    doc.add_paragraph(
        "Implementering: Naturlig parfymkomposition vid 0.3–0.5 % i slutprodukt, "
        "alternativt färdig komposition från Givaudan, Firmenich eller IFF med "
        '"natural origin"-certifiering.'
    )
    doc.add_paragraph(
        "Allergendeklaration: Vid >0.01 % i rinse-off-produkter ska parfymallergener "
        "deklareras (limonene, linalool etc.) enligt EU 1223/2009."
    )

    doc.add_page_break()

    # ── 3. Shampoo ──
    doc.add_heading("3. Formulering 1: SCHAMPO", level=1)

    doc.add_heading("Designprincip", level=2)
    doc.add_paragraph(
        "Baserat på analys av 362 naturliga schampon: det bäst presterande sulfatfria systemet "
        "använder Sodium Cocoyl Isethionate som mild primär tensid med Cocamidopropyl Betaine "
        "som co-surfaktant och en glukosidtensid för extra mildhet. Kombinationen ger rik lödder, "
        "god rengöring och minimal irritation."
    )

    doc.add_heading("INCI-formel", level=2)
    _add_table(doc,
        ["#", "INCI-namn", "Andel (%)", "Funktion", "Fas"],
        [
            ["1", "Aqua", "till 100%", "Lösningsmedel", "A"],
            ["2", "Sodium Cocoyl Isethionate", "12.0", "Primär anjonisk tensid (mild, rik lödder)", "B"],
            ["3", "Cocamidopropyl Betaine", "8.0", "Co-surfaktant (amfoter, skumförstärkare)", "A"],
            ["4", "Coco-Glucoside", "4.0", "Nonijonisk co-surfaktant (glukosidbaserad)", "A"],
            ["5", "Glycerin", "3.0", "Humektant", "A"],
            ["6", "Aloe Barbadensis Leaf Juice", "2.0", "Lugnande, fuktgivande", "A"],
            ["7", "Panthenol", "1.0", "Provitamin B5, stärker hårstråna", "A"],
            ["8", "Guar Hydroxypropyltrimonium Chloride", "0.30", "Konditionerande polymer", "C"],
            ["9", "Niacinamide", "0.30", "Vitamin B3, hårfolliklar, sebumbalans", "C"],
            ["10", "Glyceryl Oleate", "0.50", "Återfettningsmedel", "B"],
            ["11", "Xanthan Gum", "0.60", "Viskositetsreglering", "C"],
            ["12", "Sodium Chloride", "0.50", "Viskositetsreglering", "C"],
            ["13", "Sodium Benzoate", "0.40", "Konservering", "C"],
            ["14", "Potassium Sorbate", "0.30", "Konservering", "C"],
            ["15", "Ethylhexylglycerin", "0.30", "Konserveringsförstärkare", "C"],
            ["16", "Sodium Phytate", "0.10", "Kelator", "C"],
            ["17", "Citric Acid", "q.s.", "pH-justering till 5.0–5.5", "C"],
            ["18", "Parfum (Natural)", "0.40", "Doft", "C"],
        ],
        col_widths=[0.8, 5.5, 2, 6, 1],
    )

    doc.add_heading("Tillverkningsinstruktioner", level=2)
    doc.add_heading("Fas A — Vattenfas (rumstemperatur → 40 °C)", level=3)
    for step in [
        "Väg upp Aqua i huvudkärlet.",
        "Tillsätt Cocamidopropyl Betaine under långsam omrörning.",
        "Tillsätt Coco-Glucoside.",
        "Tillsätt Glycerin och Aloe Barbadensis Leaf Juice.",
        "Rör tills homogent.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Fas B — Smältfas (65–70 °C)", level=3)
    for step in [
        "Smält Sodium Cocoyl Isethionate (SCI) i separat kärl vid 65–70 °C.",
        "Tillsätt Glyceryl Oleate till smältan.",
        "Tillsätt fas B sakta till fas A under omrörning.",
        "Rör tills SCI är helt upplöst.",
        "Låt svalna till under 40 °C.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Fas C — Kallblandning (under 40 °C)", level=3)
    for step in [
        "Tillsätt Panthenol, Niacinamide, Guar Hydroxypropyltrimonium Chloride.",
        "Tillsätt Xanthan Gum (fördispergera i glycerin vid behov).",
        "Tillsätt Sodium Chloride för viskositetsjustering.",
        "Tillsätt konserveringssystemet.",
        "Tillsätt Sodium Phytate och Parfum.",
        "Justera pH med Citric Acid till 5.0–5.5.",
        "Rör långsamt 10–15 minuter, undvik luftinslagning.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Specifikationer", level=2)
    _add_table(doc,
        ["Parameter", "Krav"],
        [
            ["pH", "5.0–5.5"],
            ["Viskositet", "3 000–6 000 cP"],
            ["Färg", "Klar till lätt opalescent"],
            ["Utseende", "Slät gel utan synliga partiklar"],
            ["Mikrobiologiskt", "Klarar EU challenge test (ISO 11930)"],
            ["Stabilitet", "3 mån accelererad (40 °C / 75 % RH) utan separation"],
        ],
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # ── 4. Conditioner ──
    doc.add_heading("4. Formulering 2: BALSAM", level=1)

    doc.add_heading("Designprincip", level=2)
    doc.add_paragraph(
        "Baserat på analys av 179 naturliga balsam: det dominerande systemet använder "
        "Cetearyl Alcohol som vaxig bas med Behentrimonium Chloride som kationisk konditionerare. "
        "Sheasmör och kokosolja är de vanligaste emolienterna. Panthenol och hydrolyserat "
        "vetemjölsprotein är de mest populära aktiva ingredienserna."
    )

    doc.add_heading("INCI-formel", level=2)
    _add_table(doc,
        ["#", "INCI-namn", "Andel (%)", "Funktion", "Fas"],
        [
            ["1", "Aqua", "till 100%", "Lösningsmedel", "A"],
            ["2", "Cetearyl Alcohol", "5.0", "Fettig co-emulgator, viskositet", "B"],
            ["3", "Behentrimonium Chloride", "2.0", "Kationisk konditionerare (primär)", "B"],
            ["4", "Glycerin", "3.0", "Humektant", "A"],
            ["5", "Butyrospermum Parkii (Shea) Butter", "2.0", "Emolient, näring", "B"],
            ["6", "Cocos Nucifera (Coconut) Oil", "1.5", "Emolient, penetrerar hårskaftet", "B"],
            ["7", "Aloe Barbadensis Leaf Juice", "2.0", "Lugnande, fuktgivande", "A"],
            ["8", "Panthenol", "1.0", "Provitamin B5, stärker hår", "C"],
            ["9", "Hydrolyzed Wheat Protein", "0.50", "Proteinreparation, volym", "C"],
            ["10", "Argania Spinosa Kernel Oil", "0.50", "Arganolja, glans", "B"],
            ["11", "Stearamidopropyl Dimethylamine", "1.0", "Sekundär konditionerare", "B"],
            ["12", "Guar Hydroxypropyltrimonium Chloride", "0.20", "Konditionerande polymer", "C"],
            ["13", "Lactic Acid", "0.30", "Cuticula-nedplattning, glans", "C"],
            ["14", "Tocopherol", "0.10", "Vitamin E, antioxidant", "B"],
            ["15", "Sodium Benzoate", "0.40", "Konservering", "C"],
            ["16", "Potassium Sorbate", "0.30", "Konservering", "C"],
            ["17", "Ethylhexylglycerin", "0.30", "Konserveringsförstärkare", "C"],
            ["18", "Sodium Phytate", "0.10", "Kelator", "C"],
            ["19", "Citric Acid", "q.s.", "pH-justering till 4.5–5.0", "C"],
            ["20", "Parfum (Natural)", "0.40", "Doft", "C"],
        ],
        col_widths=[0.8, 5.5, 2, 6, 1],
    )

    doc.add_heading("Tillverkningsinstruktioner", level=2)
    doc.add_heading("Fas A — Vattenfas (70–75 °C)", level=3)
    for step in [
        "Väg upp Aqua, Glycerin, Aloe Barbadensis Leaf Juice i huvudkärlet.",
        "Värm till 70–75 °C under omrörning.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Fas B — Oljefas (70–75 °C)", level=3)
    for step in [
        "I separat kärl: smält Cetearyl Alcohol vid ~70 °C.",
        "Tillsätt Behentrimonium Chloride, Stearamidopropyl Dimethylamine, Shea Butter, Coconut Oil, Argan Oil.",
        "Tillsätt Tocopherol.",
        "Rör tills allt är smält och homogent.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Emulgering", level=3)
    for step in [
        "Häll fas B sakta i fas A under kraftig omrörning (high-shear mixer).",
        "Emulgera i 3–5 minuter tills emulsionen är slät och homogen.",
        "Fortsätt omrörning under avkylning till ~40 °C.",
    ]:
        doc.add_paragraph(step, style="List Number")

    p = doc.add_paragraph()
    r = p.add_run("OBS: ")
    r.bold = True
    p.add_run("För snabb avkylning ger kornig textur. Låt svalna gradvis.")

    doc.add_heading("Fas C — Kallblandning (under 40 °C)", level=3)
    for step in [
        "Tillsätt Panthenol, Hydrolyzed Wheat Protein, Guar Hydroxypropyltrimonium Chloride.",
        "Tillsätt konserveringssystemet.",
        "Tillsätt Sodium Phytate och Parfum.",
        "Justera pH med Citric Acid / Lactic Acid till 4.5–5.0.",
        "Rör långsamt ytterligare 15–20 minuter.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Specifikationer", level=2)
    _add_table(doc,
        ["Parameter", "Krav"],
        [
            ["pH", "4.5–5.0"],
            ["Viskositet", "15 000–30 000 cP"],
            ["Färg", "Vit till off-white"],
            ["Utseende", "Slät krämig emulsion utan korn/klumpar"],
            ["Mikrobiologiskt", "Klarar EU challenge test (ISO 11930)"],
            ["Stabilitet", "3 mån accelererad utan separation eller viskositetsförlust"],
        ],
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # ── 5. Body Wash ──
    doc.add_heading("5. Formulering 3: BODY WASH", level=1)

    doc.add_heading("Designprincip", level=2)
    doc.add_paragraph(
        "Baserat på analys av 145 body wash-produkter: de mildaste och bäst rankade använder "
        "Cocamidopropyl Betaine som primär tensid, kompletterad med Coco-Glucoside och "
        "Decyl Glucoside för den mildaste möjliga rengöringsprofilen. Högre andel glycerin (5 %) "
        "och aloe vera ger fuktbalans. Inga sulfater."
    )

    doc.add_heading("INCI-formel", level=2)
    _add_table(doc,
        ["#", "INCI-namn", "Andel (%)", "Funktion", "Fas"],
        [
            ["1", "Aqua", "till 100%", "Lösningsmedel", "A"],
            ["2", "Cocamidopropyl Betaine", "10.0", "Primär amfoter tensid (extremt mild)", "A"],
            ["3", "Coco-Glucoside", "5.0", "Co-surfaktant (nonijonisk, botanisk)", "A"],
            ["4", "Decyl Glucoside", "3.0", "Co-surfaktant (extra mildhet)", "A"],
            ["5", "Glycerin", "5.0", "Humektant (högre andel för hud)", "A"],
            ["6", "Aloe Barbadensis Leaf Juice", "2.0", "Lugnande, antiinflammatorisk", "A"],
            ["7", "Sodium Lauroyl Sarcosinate", "2.0", "Mild anjonisk tensid (lödder)", "A"],
            ["8", "Panthenol", "0.50", "Provitamin B5, hudbarriärstöd", "B"],
            ["9", "Niacinamide", "0.30", "Vitamin B3, hudfunktion", "B"],
            ["10", "Chamomilla Recutita Flower Extract", "0.30", "Kamomill, lugnande", "B"],
            ["11", "Glyceryl Oleate", "0.50", "Återfettningsmedel", "A"],
            ["12", "Sodium Cocoamphoacetate", "2.0", "Amfoter co-surfaktant", "A"],
            ["13", "Hydroxypropyl Methylcellulose", "0.40", "Viskositetsreglering", "B"],
            ["14", "Xanthan Gum", "0.30", "Viskositetsreglering", "B"],
            ["15", "Sodium Benzoate", "0.40", "Konservering", "B"],
            ["16", "Potassium Sorbate", "0.30", "Konservering", "B"],
            ["17", "Ethylhexylglycerin", "0.30", "Konserveringsförstärkare", "B"],
            ["18", "Sodium Phytate", "0.10", "Kelator", "B"],
            ["19", "Citric Acid", "q.s.", "pH-justering till 5.0–5.5", "B"],
            ["20", "Parfum (Natural)", "0.40", "Doft", "B"],
        ],
        col_widths=[0.8, 5.5, 2, 6, 1],
    )

    doc.add_heading("Tillverkningsinstruktioner", level=2)
    doc.add_heading("Fas A — Huvudfas (rumstemperatur → 35 °C)", level=3)
    for step in [
        "Väg upp Aqua i huvudkärlet.",
        "Tillsätt Cocamidopropyl Betaine långsamt under mild omrörning.",
        "Tillsätt Coco-Glucoside och Decyl Glucoside.",
        "Tillsätt Sodium Lauroyl Sarcosinate och Sodium Cocoamphoacetate.",
        "Tillsätt Glycerin, Aloe Barbadensis Leaf Juice, Glyceryl Oleate.",
        "Rör tills homogent. Undvik luftinslagning.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Fas B — Tillsatser (under 35 °C)", level=3)
    for step in [
        "Fördispergera HPMC och Xanthan Gum i en liten mängd glycerin, tillsätt sedan till batch.",
        "Tillsätt Panthenol, Niacinamide, Chamomilla Recutita-extrakt.",
        "Tillsätt konserveringssystemet.",
        "Tillsätt Sodium Phytate och Parfum.",
        "Justera pH med Citric Acid till 5.0–5.5.",
        "Rör långsamt 10–15 minuter.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Specifikationer", level=2)
    _add_table(doc,
        ["Parameter", "Krav"],
        [
            ["pH", "5.0–5.5"],
            ["Viskositet", "2 000–5 000 cP"],
            ["Färg", "Klar till lätt opalescent"],
            ["Utseende", "Slät gel"],
            ["Mikrobiologiskt", "Klarar EU challenge test (ISO 11930)"],
            ["Stabilitet", "3 mån accelererad utan separation"],
        ],
        col_widths=[4, 12],
    )

    doc.add_page_break()

    # ── 6. Shared ingredient overview ──
    doc.add_heading("6. Gemensam ingrediensöversikt", level=1)
    doc.add_paragraph("Ingredienser som alla tre produkter delar, vilket skapar en sammanhållen produktfamilj:")
    _add_table(doc,
        ["Ingrediens", "Schampo", "Balsam", "Body Wash", "Syfte"],
        [
            ["Aqua", "bas", "bas", "bas", "Lösningsmedel"],
            ["Glycerin", "3%", "3%", "5%", "Fuktbevarare"],
            ["Aloe Barbadensis Leaf Juice", "2%", "2%", "2%", "Lugnande aktiv"],
            ["Panthenol", "1%", "1%", "0.5%", "Provitamin B5"],
            ["Sodium Benzoate", "0.4%", "0.4%", "0.4%", "Konservering"],
            ["Potassium Sorbate", "0.3%", "0.3%", "0.3%", "Konservering"],
            ["Ethylhexylglycerin", "0.3%", "0.3%", "0.3%", "Konserveringsförstärkare"],
            ["Sodium Phytate", "0.1%", "0.1%", "0.1%", "Kelator"],
            ["Parfum (Natural)", "0.4%", "0.4%", "0.4%", "Doft"],
        ],
        col_widths=[5, 2, 2, 2, 4],
    )

    # ── 7. Suppliers ──
    doc.add_heading("7. Råvaruleverantörer (förslag)", level=1)
    _add_table(doc,
        ["Råvara", "Leverantörsförslag"],
        [
            ["Sodium Cocoyl Isethionate", "BASF (Plantapon SCI), Innospec"],
            ["Cocamidopropyl Betaine", "Evonik (Tego Betain), Solvay"],
            ["Coco-Glucoside / Decyl Glucoside", "BASF (Plantacare), Galaxy Surfactants"],
            ["Behentrimonium Chloride", "Evonik (Varisoft BT 85), Clariant"],
            ["Cetearyl Alcohol", "BASF, KLK Oleo"],
            ["Shea Butter", "AAK (Lipex SheaSoft)"],
            ["Konserveringssystem", "Lonza (Geogard Ultra), Ashland"],
            ["Parfym", "Givaudan, Firmenich, IFF"],
        ],
        col_widths=[6, 10],
    )

    doc.add_page_break()

    # ── 8. Testing ──
    doc.add_heading("8. Testprotokoll", level=1)

    doc.add_heading("Fas 1: Labbskala (1–2 kg batcher)", level=2)
    for s in [
        "Formulera alla tre produkter.",
        "Mät pH, viskositet, utseende.",
        "Grundläggande skumtest (cylinder-metod).",
        "Sensorisk bedömning (lukt, textur, löddring, sköljbarhet).",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Fas 2: Stabilitetstester", level=2)
    for s in [
        "Centrifugtest: 3 000 rpm, 30 min — ingen separation tillåten.",
        "Frys-tö-cykling: 5 cykler (−10 °C / +40 °C, 24 h vardera).",
        "Accelererad åldring: 40 °C / 75 % RH i 3 månader.",
        "Rumstemperatur: 25 °C i 6 månader (långtidsstabilitet).",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Fas 3: Mikrobiologisk testning", level=2)
    for s in [
        "Challenge test (ISO 11930 / EP Method A) — obligatorisk för EU-marknaden.",
        "Total aerob mikrobiell räkning vid batch-frisläpp.",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Fas 4: Användartest", level=2)
    for s in [
        "Paneltest med 20–30 personer (unisex, varierande hårtyper).",
        "Bedöm: skumkvalitet, rengöring, hårkänsla efter tork, hudkänsla, doft.",
        "Jämförelsetest mot ledande naturliga märken (Maria Nila, Sachajuan, Davines).",
    ]:
        doc.add_paragraph(s, style="List Number")

    # ── 9. Regulatory ──
    doc.add_heading("9. Regulatoriskt", level=1)

    doc.add_heading("EU Cosmetic Regulation 1223/2009", level=2)
    for s in [
        "Alla ingredienser ska vara listade i EU CosIng-databasen.",
        "CPNP-notifikation krävs innan produkten säljs.",
        "Responsible Person (RP) inom EU ska utses.",
        "Product Information File (PIF) med säkerhetsbedömning krävs.",
        "Alla parfymallergener ovan 0.01 % i rinse-off ska deklareras.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Claims", level=2)
    doc.add_paragraph("Med dessa formuleringar kan följande påståenden göras:")
    for s in [
        '"Sulfatfri" / "Sulfate-free"',
        '"Silikonfri" / "Silicone-free"',
        '"Parabensfri" / "Paraben-free"',
        '"Vegansk" (verifiera med leverantör att inga animaliska råvaror)',
        '"Med naturliga ingredienser"',
        '"Mild / Gentle"',
    ]:
        doc.add_paragraph(s, style="List Bullet")

    p = doc.add_paragraph()
    r = p.add_run('Undvik: ')
    r.bold = True
    p.add_run('"Naturlig" som ensamt claim utan certifiering — kan vara vilseledande.')

    # ── 10. Cost ──
    doc.add_heading("10. Kostnadskalkyl (uppskattning)", level=1)
    _add_table(doc,
        ["Komponent", "Schampo", "Balsam", "Body Wash"],
        [
            ["Råvaror (250 ml)", "8–12 SEK", "10–15 SEK", "8–12 SEK"],
            ["Förpackning (flaska + lock)", "5–8 SEK", "5–8 SEK", "5–8 SEK"],
            ["Tillverkning + fyllning", "3–5 SEK", "3–5 SEK", "3–5 SEK"],
            ["Total per produkt", "16–25 SEK", "18–28 SEK", "16–25 SEK"],
            ["Paket (3 st)", "", "50–78 SEK", ""],
        ],
        col_widths=[5, 3.5, 3.5, 3.5],
    )

    doc.add_paragraph(
        "Vid konsumentpris 399 SEK och antagen 50 % rabatt till klubbar (inköp 200 SEK): "
        "bruttomarginal 60–75 % vid kontraktstillverkning i större volymer."
    )
    doc.add_paragraph(
        "Priserna baseras på nordeuropeiska leverantörspriser vid MOQ 500–1 000 kg. "
        "Skalas ner vid högre volymer."
    )

    # ── Save ──
    path = "Roots_Formuleringar_Kemistspec.docx"
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build_document()
